"""
Insert Ramachandran plot figure (PNG) into the docx manuscript
after the Ramachandran results body paragraph, as Figure 3.
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy, io

PATH     = '/sessions/confident-keen-gauss/mnt/First Iteration/revision-notes-0-02052026045611.docx'
IMG_PATH = '/sessions/confident-keen-gauss/mnt/outputs/ramachandran_plot.png'

doc = Document(PATH)

# ── Locate anchor paragraph (Ramachandran body text) ─────────────────────────
anchor = None
for para in doc.paragraphs:
    if para.text.startswith('Ramachandran plot analysis was run'):
        anchor = para
        break

if anchor is None:
    raise RuntimeError("Anchor paragraph not found")

print("Anchor: " + anchor.text[:70])

# ── Build image paragraph ─────────────────────────────────────────────────────
# We must insert using the docx inline API, then move the resulting <w:p> element
# to the correct position.

# Step 1: temporarily append image paragraph at end of document body
img_para = doc.add_paragraph()
img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = img_para.add_run()
run.add_picture(IMG_PATH, width=Inches(5.8))

# Step 2: build caption paragraph
cap_para = doc.add_paragraph()
cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap_run = cap_para.add_run(
    "Figure 3. Ramachandran plot analysis of CCR5 (PDB ID: 4MBS, left panel) "
    "and CCL5 (PDB ID: 1RTO, right panel). Circles represent residues in the "
    "favored region (blue), squares in the allowed region (blue), and triangles "
    "mark disallowed outliers (red). Shaded background zones indicate the "
    "classically defined favored (dark blue) and broadly allowed (light blue) "
    "regions of the Ramachandran map. Both structures satisfy the <15% "
    "disallowed threshold; all outlying residues correspond to glycines, whose "
    "lack of a β-carbon renders them inherently more conformationally flexible."
)
cap_run.italic = True
cap_run.font.size = Pt(9)

# Step 3: move both new paragraphs to just after the anchor
body = doc.element.body
body.remove(img_para._p)
body.remove(cap_para._p)

anchor._p.addnext(cap_para._p)
anchor._p.addnext(img_para._p)

doc.save(PATH)
print("Figure 3 inserted and saved.")

# Verify
doc2 = Document(PATH)
for i, para in enumerate(doc2.paragraphs):
    if 'Figure 3' in para.text or 'Ramachandran plot analysis' in para.text:
        print("  Para[" + str(i) + "]: " + para.text[:90])
