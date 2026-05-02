from docx import Document

PATH = '/sessions/confident-keen-gauss/mnt/First Iteration/revision-notes-0-02052026045611.docx'
doc = Document(PATH)

# ── Para [16]: add PDB ID 1RTO to CCL5 entry ─────────────────────────────────
OLD16 = ("and CCL5 (ligand for protein–protein docking) were retrieved from "
         "the RCSB Protein Data Bank (https://www.rcsb.org) in PDB format.")

NEW16 = ("and CCL5 (ligand for protein–protein docking; PDB ID: 1RTO; "
         "https://www.rcsb.org/structure/1RTO) were retrieved from "
         "the RCSB Protein Data Bank (https://www.rcsb.org) in PDB format.")

# ── Para [30]: first CCL5 mention in results ──────────────────────────────────
OLD30 = "Protein–protein docking of CCL5 to CCR5 via HDOCK yielded 10 ranked models (Table 2)."
NEW30 = "Protein–protein docking of CCL5 (PDB ID: 1RTO) to CCR5 (PDB ID: 4MBS) via HDOCK yielded 10 ranked models (Table 2)."

for para_idx, old, new in [(16, OLD16, NEW16), (30, OLD30, NEW30)]:
    para = doc.paragraphs[para_idx]
    patched = False
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            print("Para[" + str(para_idx) + "] patched OK")
            patched = True
            break
    if not patched:
        print("Para[" + str(para_idx) + "] WARNING: pattern not found")

doc.save(PATH)
print("Saved.")

# Verify
doc2 = Document(PATH)
for idx in [16, 30]:
    print("\nPara[" + str(idx) + "]:\n" + doc2.paragraphs[idx].text[:400])
