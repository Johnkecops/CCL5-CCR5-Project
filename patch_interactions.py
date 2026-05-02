"""
Insert 2D/3D interaction analysis into manuscript:
  Methods: after Para[21] (re-docking validation), before Para[22]
  Results: after Para[37] (docking affinity paragraph), before Table 1
           - heading + body + 2D figure + caption + 3D figure + caption + table heading
"""

from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import copy, os

DOC_PATH  = '/sessions/confident-keen-gauss/mnt/First Iteration/revision-notes-0-02052026045611.docx'
IMG_2D    = '/sessions/confident-keen-gauss/mnt/First Iteration/interaction_2D.jpg'
IMG_3D    = '/sessions/confident-keen-gauss/mnt/First Iteration/interaction_3D.jpg'

doc = Document(DOC_PATH)

# ── Helpers ──────────────────────────────────────────────────────────────────
def clone_normal_para(ref_para, text):
    new_p = OxmlElement('w:p')
    pPr = ref_para._p.find(qn('w:pPr'))
    if pPr is not None: new_p.append(copy.deepcopy(pPr))
    r = OxmlElement('w:r')
    first_r = ref_para._p.find(qn('w:r'))
    if first_r is not None:
        rPr = first_r.find(qn('w:rPr'))
        if rPr is not None: r.append(copy.deepcopy(rPr))
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t); new_p.append(r)
    return new_p

def make_bold_para(ref_para, text):
    new_p = OxmlElement('w:p')
    pPr = ref_para._p.find(qn('w:pPr'))
    if pPr is not None: new_p.append(copy.deepcopy(pPr))
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rPr.append(OxmlElement('w:b'))
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t); new_p.append(r)
    return new_p

def make_italic_para(ref_para, text):
    new_p = OxmlElement('w:p')
    pPr = ref_para._p.find(qn('w:pPr'))
    if pPr is not None: new_p.append(copy.deepcopy(pPr))
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rPr.append(OxmlElement('w:i'))
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t); new_p.append(r)
    return new_p

def make_image_para(doc, image_path, width_in=5.5):
    para = doc.add_paragraph()
    para.add_run().add_picture(image_path, width=Inches(width_in))
    img_p = copy.deepcopy(para._p)
    para._p.getparent().remove(para._p)
    return img_p

# ── Text content ─────────────────────────────────────────────────────────────

METHODS_TEXT = (
    "Intermolecular interactions between maraviroc and CCR5 were characterised using a "
    "custom Python pipeline built on Biopython and RDKit. Protein and ligand heavy atoms were "
    "parsed from the 4MBS crystal structure. Residues with any heavy atom within 5.5 Angstroms "
    "of the ligand were identified as contact residues. Interactions were then assigned to one "
    "of five categories: hydrogen bonds (donor-acceptor distance below 3.5 Angstroms, involving "
    "N, O, F, or S atoms), salt bridges (charged N or O atoms of Arg/Lys/His and Asp/Glu "
    "residues within 5.5 Angstroms), pi interactions (aromatic residues within 5.5 Angstroms), "
    "hydrophobic contacts (C-C pairs within 5.0 Angstroms), and van der Waals contacts (all "
    "remaining heavy-atom pairs within 4.0 Angstroms). Where multiple interaction types were "
    "possible for a single residue, the highest-priority category was assigned in the order "
    "listed. Two-dimensional interaction diagrams were generated with RDKit 2D coordinates and "
    "annotated with residue labels and interaction lines colour-coded by type. Three-dimensional "
    "binding pocket visualisations were produced with Matplotlib, showing the CCR5 C-alpha "
    "backbone within 14 Angstroms of the ligand centroid together with interacting residues "
    "and colour-coded contact lines."
)

RESULTS_HEADING = "Intermolecular Interactions of Maraviroc within the CCR5 Binding Pocket"

RESULTS_BODY = (
    "Interaction analysis of the maraviroc-CCR5 complex (PDB: 4MBS) identified 22 contact "
    "residues within 5.5 Angstroms of the ligand (Table 3). Five residues formed direct "
    "hydrogen bonds: GLU283 (2.80 Angstroms, N2-OE1), TYR251 (3.02 Angstroms, N5-OH), "
    "TYR37 (3.11 Angstroms, N3-OH), THR259 (3.42 Angstroms, F2-OG1), and THR195 "
    "(3.46 Angstroms, F2-OG1). Nine residues contributed pi or aromatic interactions, "
    "including PHE182, PHE109, TRP86, TYR108, TYR89, TRP248, PHE112, TRP190, and PHE79 "
    "(3.21-5.32 Angstroms). Eight residues formed hydrophobic contacts: LEU255, ILE198, "
    "MET279, MET287, GLN194, THR284, LYS191, and THR105 (3.72-4.84 Angstroms). "
    "The 2D interaction diagram (Figure 5) maps these contacts around the ligand scaffold, "
    "while the 3D binding pocket view (Figure 6) shows that maraviroc sits within a well-defined "
    "hydrophobic cavity lined predominantly by aromatic and aliphatic residues, with hydrogen "
    "bond anchors contributed by GLU283 and the tyrosine/threonine cluster. This binding "
    "geometry is consistent with published crystallographic analyses of the 4MBS structure "
    "(Tan et al., 2013)."
)

CAP_2D = (
    "Figure 5. Two-dimensional interaction diagram of maraviroc (MRV) within the CCR5 binding "
    "pocket (PDB: 4MBS). Residue boxes are colour-coded by interaction type: blue = hydrogen "
    "bond, orange = hydrophobic contact, purple = pi interaction, red = salt bridge, grey = "
    "van der Waals. Distances in Angstroms."
)

CAP_3D = (
    "Figure 6. Three-dimensional view of maraviroc within the CCR5 binding pocket (PDB: 4MBS). "
    "The CCR5 C-alpha trace (blue) is shown within 14 Angstroms of the ligand centroid. "
    "Interacting residues are plotted as spheres colour-coded by interaction type. "
    "Coloured dashed/dotted lines indicate specific contacts."
)

# Interaction table data
TABLE_ROWS = [
    ("GLU283", "A", "283", "H-bond",        "2.80", "N2",  "OE1"),
    ("TYR251", "A", "251", "H-bond",        "3.02", "N5",  "OH"),
    ("TYR37",  "A", "37",  "H-bond",        "3.11", "N3",  "OH"),
    ("PHE182", "A", "182", "Pi interaction","3.21", "F1",  "CE2"),
    ("THR259", "A", "259", "H-bond",        "3.42", "F2",  "OG1"),
    ("THR195", "A", "195", "H-bond",        "3.46", "F2",  "OG1"),
    ("PHE109", "A", "109", "Pi interaction","3.48", "C25", "CE1"),
    ("TRP86",  "A", "86",  "Pi interaction","3.51", "N4",  "CB"),
    ("TYR108", "A", "108", "Pi interaction","3.62", "C20", "OH"),
    ("TYR89",  "A", "89",  "Pi interaction","3.64", "C14", "CD2"),
    ("TRP248", "A", "248", "Pi interaction","3.66", "C28", "CZ3"),
    ("PHE112", "A", "112", "Pi interaction","3.71", "C29", "CG"),
    ("LEU255", "A", "255", "Hydrophobic",   "3.72", "C23", "CD1"),
    ("GLN194", "A", "194", "Hydrophobic",   "3.84", "C22", "CB"),
    ("ILE198", "A", "198", "Hydrophobic",   "3.84", "C22", "CD1"),
    ("MET279", "A", "279", "Hydrophobic",   "3.87", "C23", "CE"),
    ("MET287", "A", "287", "Hydrophobic",   "3.93", "C20", "CG"),
    ("THR284", "A", "284", "Hydrophobic",   "4.06", "C19", "CA"),
    ("LYS191", "A", "191", "Hydrophobic",   "4.58", "C18", "CA"),
    ("THR105", "A", "105", "Hydrophobic",   "4.84", "C9",  "CG2"),
    ("TRP190", "A", "190", "Pi interaction","5.25", "F1",  "O"),
    ("PHE79",  "A", "79",  "Pi interaction","5.32", "C29", "CE2"),
]

TABLE_CAPTION = (
    "Table 3. Amino acid residues of CCR5 involved in interactions with maraviroc "
    "(PDB: 4MBS), classified by interaction type. Distances computed over heavy atoms."
)

# ── Find anchor paragraphs ───────────────────────────────────────────────────
anchor_methods = None   # Para[21] — insert AFTER
anchor_ppi     = None   # Para[22] Protein-Protein heading — verify order

# Results anchor: Para[38] "Table 1..." — insert BEFORE
anchor_table1  = None

ref_para = None

for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if 'native ligand of CCR5 was re-docked' in txt:
        anchor_methods = p
        print(f"[{i}] anchor_methods")
    if txt.startswith('Protein–Protein Molecular Docking'):
        anchor_ppi = p
        print(f"[{i}] anchor_ppi")
    if 'Table 1' in txt and 'Binding results' in txt:
        anchor_table1 = p
        print(f"[{i}] anchor_table1")
    if 'Ramachandran plot analysis was run' in txt and ref_para is None:
        ref_para = p

assert anchor_methods, "anchor_methods not found"
assert anchor_table1, "anchor_table1 not found"
assert ref_para, "ref_para not found"

# ── Methods: insert after anchor_methods (chain: methods → ppi) ─────────────
meth_p = clone_normal_para(ref_para, METHODS_TEXT)
anchor_methods._p.addnext(meth_p)
print("Methods paragraph inserted.")

# ── Results: insert before Table 1 ──────────────────────────────────────────
# Build all paragraphs and a DOCX table
# Order: heading → body → 2d_img → cap_2d → 3d_img → cap_3d → table_caption → table

# Build image paragraphs
img_2d_p = make_image_para(doc, IMG_2D, 5.8)
img_3d_p = make_image_para(doc, IMG_3D, 5.8)

heading_p  = make_bold_para(ref_para, RESULTS_HEADING)
body_p     = clone_normal_para(ref_para, RESULTS_BODY)
cap_2d_p   = make_italic_para(ref_para, CAP_2D)
cap_3d_p   = make_italic_para(ref_para, CAP_3D)
tab_cap_p  = make_italic_para(ref_para, TABLE_CAPTION)

# Build the interaction table
from docx.shared import Pt
from docx.oxml.ns import qn as _qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

tbl = doc.add_table(rows=1, cols=7)
tbl.style = 'Table Grid'

# Header row
hdr_cells = tbl.rows[0].cells
headers = ["Residue", "Chain", "Seq #", "Interaction Type",
           "Distance (Å)", "Ligand Atom", "Protein Atom"]
for j, h in enumerate(headers):
    hdr_cells[j].text = h
    run = hdr_cells[j].paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(8)

# Data rows
for row_data in TABLE_ROWS:
    row_cells = tbl.add_row().cells
    for j, val in enumerate(row_data):
        row_cells[j].text = val
        row_cells[j].paragraphs[0].runs[0].font.size = Pt(8)

tbl_p = copy.deepcopy(tbl._tbl)
doc.tables[-1]._tbl.getparent().remove(doc.tables[-1]._tbl)

# Insert in order (chained addnext from heading):
anchor_table1._p.addprevious(heading_p)
heading_p.addnext(body_p)
body_p.addnext(img_2d_p)
img_2d_p.addnext(cap_2d_p)
cap_2d_p.addnext(img_3d_p)
img_3d_p.addnext(cap_3d_p)
cap_3d_p.addnext(tab_cap_p)
tab_cap_p.addnext(tbl_p)

print("Results section inserted.")

doc.save(DOC_PATH)
print(f"Saved: {DOC_PATH}")

# ── Verify ───────────────────────────────────────────────────────────────────
doc_v = Document(DOC_PATH)
print("\n--- Verification ---")
for i, p in enumerate(doc_v.paragraphs):
    txt = p.text.strip()
    if txt and any(k in txt for k in ['Intermolecular','Re-docked','Re-docking','maraviroc'
                                       'Table 1','Table 3','Figure 5','Figure 6',
                                       'AutoDock Vina generated','custom Python pipeline']):
        print(f"[{i:3d}] {txt[:90]}")
