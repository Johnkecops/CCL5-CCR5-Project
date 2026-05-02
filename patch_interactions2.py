"""
Insert 2D/3D interaction analysis - final version.
Methods paragraph is already inserted (Para[22]).
This script inserts Results section before Table 1.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

DOC_PATH = '/sessions/confident-keen-gauss/mnt/First Iteration/revision-notes-0-02052026045611.docx'
IMG_2D   = '/sessions/confident-keen-gauss/mnt/First Iteration/interaction_2D.jpg'
IMG_3D   = '/sessions/confident-keen-gauss/mnt/First Iteration/interaction_3D.jpg'

doc = Document(DOC_PATH)

# Verify methods paragraph already exists
for i, p in enumerate(doc.paragraphs):
    if 'custom Python pipeline' in p.text:
        print(f"[{i}] Methods interaction para already present: OK")
        break
else:
    print("WARNING: methods para not found")

# ── Helpers ──────────────────────────────────────────────────────────────────
def clone_para(ref_para, text):
    new_p = OxmlElement('w:p')
    pPr = ref_para._p.find(qn('w:pPr'))
    if pPr is not None: new_p.append(copy.deepcopy(pPr))
    r = OxmlElement('w:r')
    first_r = ref_para._p.find(qn('w:r'))
    if first_r is not None:
        rPr = first_r.find(qn('w:rPr'))
        if rPr is not None: r.append(copy.deepcopy(rPr))
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t); new_p.append(r)
    return new_p

def bold_para(ref_para, text):
    new_p = OxmlElement('w:p')
    pPr = ref_para._p.find(qn('w:pPr'))
    if pPr is not None: new_p.append(copy.deepcopy(pPr))
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr'); rPr.append(OxmlElement('w:b'))
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t); new_p.append(r)
    return new_p

def italic_para(ref_para, text):
    new_p = OxmlElement('w:p')
    pPr = ref_para._p.find(qn('w:pPr'))
    if pPr is not None: new_p.append(copy.deepcopy(pPr))
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr'); rPr.append(OxmlElement('w:i'))
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t); new_p.append(r)
    return new_p

def image_para(doc, path, width_in=5.6):
    para = doc.add_paragraph()
    para.add_run().add_picture(path, width=Inches(width_in))
    p = copy.deepcopy(para._p)
    para._p.getparent().remove(para._p)
    return p

def set_table_borders(table):
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for bname in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{bname}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        tblBorders.append(b)
    tblPr.append(tblBorders)

# ── Text ─────────────────────────────────────────────────────────────────────
RESULTS_HEAD = "Intermolecular Interactions of Maraviroc within the CCR5 Binding Pocket"

RESULTS_BODY = (
    "Interaction analysis of the maraviroc-CCR5 complex (PDB: 4MBS) identified 22 contact "
    "residues within 5.5 Angstroms of the ligand (Table 3). Five residues formed direct "
    "hydrogen bonds: GLU283 (2.80 Angstroms, N2-OE1), TYR251 (3.02 Angstroms, N5-OH), "
    "TYR37 (3.11 Angstroms, N3-OH), THR259 (3.42 Angstroms, F2-OG1), and THR195 "
    "(3.46 Angstroms, F2-OG1). Nine residues contributed pi or aromatic interactions, "
    "including PHE182, PHE109, TRP86, TYR108, TYR89, TRP248, PHE112, TRP190, and PHE79 "
    "(3.21-5.32 Angstroms). Eight residues formed hydrophobic contacts: LEU255, ILE198, "
    "MET279, MET287, GLN194, THR284, LYS191, and THR105 (3.72-4.84 Angstroms). "
    "The 2D interaction diagram (Figure 5) maps these contacts around the ligand scaffold, "
    "while the 3D binding pocket view (Figure 6) shows that maraviroc occupies a well-defined "
    "hydrophobic cavity lined by aromatic and aliphatic residues, with hydrogen bond anchors "
    "from GLU283 and the tyrosine/threonine cluster. This binding geometry agrees with "
    "published crystallographic analyses of the 4MBS structure (Tan et al., 2013)."
)

CAP_2D = (
    "Figure 5. Two-dimensional interaction diagram of maraviroc within the CCR5 binding "
    "pocket (PDB: 4MBS). Residue labels are colour-coded by interaction type: blue = hydrogen "
    "bond, orange = hydrophobic contact, purple = pi interaction, grey = van der Waals. "
    "Distances are in Angstroms."
)

CAP_3D = (
    "Figure 6. Three-dimensional view of the maraviroc-CCR5 binding pocket (PDB: 4MBS). "
    "The CCR5 C-alpha trace (blue) is shown within 14 Angstroms of the ligand centroid. "
    "Interacting residues are displayed as colour-coded spheres. Dashed and dotted lines "
    "denote specific contacts by interaction type."
)

TABLE_CAPTION = (
    "Table 3. Amino acid residues of CCR5 involved in intermolecular interactions with "
    "maraviroc (PDB: 4MBS). Distances are computed over heavy atoms. Interactions assigned "
    "by priority: H-bond > salt bridge > pi > hydrophobic > van der Waals."
)

TABLE_ROWS = [
    ("GLU283","A","283","H-bond","2.80","N2","OE1"),
    ("TYR251","A","251","H-bond","3.02","N5","OH"),
    ("TYR37", "A","37", "H-bond","3.11","N3","OH"),
    ("PHE182","A","182","Pi interaction","3.21","F1","CE2"),
    ("THR259","A","259","H-bond","3.42","F2","OG1"),
    ("THR195","A","195","H-bond","3.46","F2","OG1"),
    ("PHE109","A","109","Pi interaction","3.48","C25","CE1"),
    ("TRP86", "A","86", "Pi interaction","3.51","N4","CB"),
    ("TYR108","A","108","Pi interaction","3.62","C20","OH"),
    ("TYR89", "A","89", "Pi interaction","3.64","C14","CD2"),
    ("TRP248","A","248","Pi interaction","3.66","C28","CZ3"),
    ("PHE112","A","112","Pi interaction","3.71","C29","CG"),
    ("LEU255","A","255","Hydrophobic","3.72","C23","CD1"),
    ("GLN194","A","194","Hydrophobic","3.84","C22","CB"),
    ("ILE198","A","198","Hydrophobic","3.84","C22","CD1"),
    ("MET279","A","279","Hydrophobic","3.87","C23","CE"),
    ("MET287","A","287","Hydrophobic","3.93","C20","CG"),
    ("THR284","A","284","Hydrophobic","4.06","C19","CA"),
    ("LYS191","A","191","Hydrophobic","4.58","C18","CA"),
    ("THR105","A","105","Hydrophobic","4.84","C9","CG2"),
    ("TRP190","A","190","Pi interaction","5.25","F1","O"),
    ("PHE79", "A","79", "Pi interaction","5.32","C29","CE2"),
]

# ── Locate anchors ────────────────────────────────────────────────────────────
ref_para = anchor_table1 = None
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if 'Ramachandran plot analysis was run' in txt and ref_para is None:
        ref_para = p
    if 'Table 1' in txt and 'Binding results' in txt:
        anchor_table1 = p
        print(f"[{i}] anchor_table1 found")

assert anchor_table1 and ref_para

# ── Build table ────────────────────────────────────────────────────────────────
tbl = doc.add_table(rows=1, cols=7)
set_table_borders(tbl)
hdr = tbl.rows[0].cells
for j, h in enumerate(["Residue","Chain","Seq #","Interaction Type",
                        "Distance (A)","Ligand Atom","Protein Atom"]):
    hdr[j].text = h
    if hdr[j].paragraphs[0].runs:
        hdr[j].paragraphs[0].runs[0].bold = True
        hdr[j].paragraphs[0].runs[0].font.size = Pt(8)
for rd in TABLE_ROWS:
    cells = tbl.add_row().cells
    for j, v in enumerate(rd):
        cells[j].text = v
        if cells[j].paragraphs[0].runs:
            cells[j].paragraphs[0].runs[0].font.size = Pt(8)

tbl_elem = copy.deepcopy(tbl._tbl)
doc.tables[-1]._tbl.getparent().remove(doc.tables[-1]._tbl)

# ── Build paragraph elements ──────────────────────────────────────────────────
h_p   = bold_para(ref_para, RESULTS_HEAD)
b_p   = clone_para(ref_para, RESULTS_BODY)
i2_p  = image_para(doc, IMG_2D, 5.6)
c2_p  = italic_para(ref_para, CAP_2D)
i3_p  = image_para(doc, IMG_3D, 5.6)
c3_p  = italic_para(ref_para, CAP_3D)
tc_p  = italic_para(ref_para, TABLE_CAPTION)

# ── Insert: heading → body → fig5 → cap5 → fig6 → cap6 → table_cap → table → Table1 ──
anchor_table1._p.addprevious(h_p)
h_p.addnext(b_p)
b_p.addnext(i2_p)
i2_p.addnext(c2_p)
c2_p.addnext(i3_p)
i3_p.addnext(c3_p)
c3_p.addnext(tc_p)
tc_p.addnext(tbl_elem)

print("All results content inserted.")
doc.save(DOC_PATH)
print(f"Saved: {DOC_PATH}")

# ── Quick verification ────────────────────────────────────────────────────────
doc_v = Document(DOC_PATH)
print("\n--- Verification ---")
for i, p in enumerate(doc_v.paragraphs):
    txt = p.text.strip()
    if txt and any(k in txt for k in ['Intermolecular','custom Python','Figure 5','Figure 6',
                                       'Table 3','Table 1','AutoDock Vina generated']):
        print(f"[{i:3d}] {txt[:85]}")
print(f"\nTotal tables in doc: {len(doc_v.tables)}")
