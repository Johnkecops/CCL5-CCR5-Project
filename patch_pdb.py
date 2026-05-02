from docx import Document

PATH = '/sessions/confident-keen-gauss/mnt/First Iteration/revision-notes-0-02052026045611.docx'
doc = Document(PATH)

# ── Para [16]: Structure Retrieval ────────────────────────────────────────────
OLD16 = ("The three-dimensional structure files of CCR5 (receptor) and CCL5 "
         "(ligand for protein–protein docking) were retrieved from the "
         "Protein Data Bank (PDB) in PDB format.")

NEW16 = ("The three-dimensional structure of CCR5 (receptor; PDB ID: 4MBS; "
         "https://www.rcsb.org/structure/4MBS) and CCL5 (ligand for "
         "protein–protein docking) were retrieved from the RCSB Protein "
         "Data Bank (https://www.rcsb.org) in PDB format.")

# ── Para [25]: Results – first CCR5 mention ───────────────────────────────────
OLD25 = "maraviroc–CCR5 interaction (Table 1)."
NEW25 = "maraviroc–CCR5 (PDB ID: 4MBS) interaction (Table 1)."

patches = [(16, OLD16, NEW25, NEW16), (25, OLD25, None, NEW25)]

for para_idx, old, _, new in [(16, OLD16, None, NEW16), (25, OLD25, None, NEW25)]:
    para = doc.paragraphs[para_idx]
    patched = False
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            print("Para[" + str(para_idx) + "] patched OK")
            patched = True
            break
    if not patched:
        print("Para[" + str(para_idx) + "] WARNING: pattern not found")

doc.save(PATH)
print("Saved.")

# Verify
doc2 = Document(PATH)
for idx in [16, 25]:
    t = doc2.paragraphs[idx].text
    print("\nPara[" + str(idx) + "]:\n" + t[:350])
