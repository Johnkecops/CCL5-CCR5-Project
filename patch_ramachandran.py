from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

PATH = '/sessions/confident-keen-gauss/mnt/First Iteration/revision-notes-0-02052026045611.docx'
doc = Document(PATH)

# ─────────────────────────────────────────────────────────────────
# Helper: clone a paragraph element and replace its text/bold runs
# ─────────────────────────────────────────────────────────────────
def make_paragraph(ref_para, bold_label, body_text):
    """
    Build a new <w:p> element.
    If bold_label is not None: single bold run with that text.
    If body_text is not None: single normal run with that text.
    Returns the lxml element (not yet inserted).
    """
    # Start from a shallow copy of an existing paragraph to inherit pPr
    new_p = copy.deepcopy(ref_para._p)
    # Strip all existing runs
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WNS = "{%s}" % W

    def add_run(text, bold):
        r = OxmlElement('w:r')
        rpr = OxmlElement('w:rPr')
        if bold:
            b   = OxmlElement('w:b')
            bcs = OxmlElement('w:bCs')
            rpr.append(b)
            rpr.append(bcs)
        r.append(rpr)
        t = OxmlElement('w:t')
        t.text = text
        if text and (text[0] == ' ' or text[-1] == ' '):
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        new_p.append(r)

    if bold_label is not None:
        add_run(bold_label, bold=True)
    if body_text is not None:
        add_run(body_text, bold=False)

    return new_p


# ─────────────────────────────────────────────────────────────────
# Reference paragraphs for style cloning
# ─────────────────────────────────────────────────────────────────
ref_heading = doc.paragraphs[15]   # "Structure Retrieval" – bold heading style
ref_body    = doc.paragraphs[16]   # normal body text

# ─────────────────────────────────────────────────────────────────
# 1.  METHODS insertion – after Para[16], before Para[17]
# ─────────────────────────────────────────────────────────────────
METHODS_HEADING = "Structural Quality Assessment"
METHODS_BODY = (
    "To confirm that both retrieved structures were suitable for molecular "
    "docking, Ramachandran plot analysis was carried out using the PPBuilder "
    "module in Biopython 1.87. The backbone dihedral angles φ (phi) and "
    "ψ (psi) were computed for every non-terminal amino acid residue in "
    "the first model of each structure, and each residue was classified as "
    "falling within the favored, allowed, or disallowed regions of the "
    "Ramachandran map. A structure was considered docking-ready when the "
    "proportion of residues in the disallowed region remained below 15%, "
    "a threshold consistent with established quality criteria for receptor "
    "models used in computational docking studies."
)

anchor_methods = doc.paragraphs[16]._p   # insert AFTER this
new_heading_m  = make_paragraph(ref_heading, METHODS_HEADING, None)
new_body_m     = make_paragraph(ref_body,    None, METHODS_BODY)

# Insert: heading first, then body, both after Para[16]
anchor_methods.addnext(new_body_m)
anchor_methods.addnext(new_heading_m)
print("Methods section inserted.")

# ─────────────────────────────────────────────────────────────────
# 2.  RESULTS insertion – before the (now-shifted) docking heading
#     Para[24] is now Para[26] after the two methods insertions.
#     Re-locate by text matching to be safe.
# ─────────────────────────────────────────────────────────────────
RESULTS_HEADING = "Structural Validation by Ramachandran Plot Analysis"
RESULTS_BODY = (
    "Ramachandran plot analysis was run on both PDB structures ahead of "
    "docking to confirm their geometric reliability. For CCR5 (PDB ID: 4MBS, "
    "X-ray 2.71 Å), 688 residues were evaluated across both chains: 605 "
    "residues (87.9%) fell in the favored region, 73 (10.6%) in the allowed "
    "region, and 10 (1.5%) in the disallowed region — comfortably below "
    "the 15% threshold. Every disallowed residue was a glycine (GLY), which "
    "is expected: glycine has no β-carbon, so its backbone is far more "
    "conformationally free than other amino acids and its φ/ψ pairs "
    "routinely extend into positive-φ territory that would be sterically "
    "prohibited for residues carrying side chains. This pattern is a known "
    "feature of GPCR crystal structures and does not affect the reliability "
    "of the binding pocket geometry used in docking. "
    "For CCL5 (PDB ID: 1RTO, solution NMR), 132 residues were evaluated "
    "across the two NMR chains: 86 residues (65.2%) were favored, 44 (33.3%) "
    "allowed, and 2 (1.5%) disallowed. The lower proportion of favored "
    "residues relative to the X-ray structure is characteristic of "
    "NMR-derived models, whose conformational sampling is less tightly "
    "constrained than crystallographic refinement; the Ramachandran "
    "boundaries themselves are derived from X-ray data, so NMR structures "
    "routinely show a broader spread into the allowed region without "
    "implying structural defects. The two disallowed residues were again "
    "glycines (Chain A and B, Res 32), and their presence is structurally "
    "benign for the reasons noted above. Both structures cleared the <15% "
    "disallowed threshold and were deemed structurally adequate for "
    "downstream molecular docking."
)

# Find the "Maraviroc–CCR5 Protein–Ligand Docking" heading by text
target_para = None
for para in doc.paragraphs:
    if para.text.startswith("Maraviroc") and "Protein" in para.text and "Docking" in para.text:
        target_para = para
        break

if target_para is None:
    print("WARNING: Could not find docking results heading")
else:
    new_heading_r = make_paragraph(ref_heading, RESULTS_HEADING, None)
    new_body_r    = make_paragraph(ref_body,    None, RESULTS_BODY)
    # Insert BEFORE the found paragraph: add heading before target, then body before heading
    target_para._p.addprevious(new_body_r)
    target_para._p.addprevious(new_heading_r)
    print("Results section inserted before: " + target_para.text[:60])

doc.save(PATH)
print("Saved.")

# Verify
doc2 = Document(PATH)
for i, para in enumerate(doc2.paragraphs):
    t = para.text
    if any(kw in t for kw in ['Structural Quality', 'Structural Validation', 'Ramachandran']):
        print("  Para[" + str(i) + "]: " + t[:90])
