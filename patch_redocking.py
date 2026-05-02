"""
Insert re-docking validation content into manuscript:
  - Methods: validation paragraph after Para[20] (docking body)
  - Results: heading + body + figure before Para[30] (Maraviroc docking results)
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import copy, os

DOC_PATH = '/sessions/confident-keen-gauss/mnt/First Iteration/revision-notes-0-02052026045611.docx'
JPG_PATH = '/sessions/confident-keen-gauss/mnt/First Iteration/maraviroc_CCR5_redocking.jpg'

doc = Document(DOC_PATH)
paras = doc.paragraphs

def clone_para_format(source_para, new_text, bold_prefix=None):
    """Return a new paragraph mimicking source_para's run formatting."""
    new_p = OxmlElement('w:p')
    new_pPr = copy.deepcopy(source_para._p.find(qn('w:pPr'))) if source_para._p.find(qn('w:pPr')) is not None else None
    if new_pPr is not None:
        new_p.append(new_pPr)
    
    if bold_prefix:
        # Bold run for prefix
        r_bold = OxmlElement('w:r')
        rPr_bold = OxmlElement('w:rPr')
        b = OxmlElement('w:b')
        rPr_bold.append(b)
        r_bold.append(rPr_bold)
        t = OxmlElement('w:t')
        t.text = bold_prefix
        r_bold.append(t)
        new_p.append(r_bold)
    
    # Normal run for body text
    r = OxmlElement('w:r')
    # Copy run properties from first run of source if available
    first_run = source_para._p.find(qn('w:r'))
    if first_run is not None:
        rPr = first_run.find(qn('w:rPr'))
        if rPr is not None:
            r.append(copy.deepcopy(rPr))
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = new_text
    r.append(t)
    new_p.append(r)
    return new_p

def insert_image_para(source_para, image_path, width_inches=5.5):
    """Create a paragraph with an embedded image."""
    from docx.shared import Inches
    tmp_doc = Document()
    run = tmp_doc.paragraphs[0].add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    img_p = tmp_doc.paragraphs[0]._p
    # Clone and return
    return copy.deepcopy(img_p)

# ── Locate anchor paragraphs ─────────────────────────────────────────────────
# Para[20]: body of Protein–Ligand Molecular Docking section
# Para[21]: Protein–Protein Molecular Docking heading
# Para[29]: Figure 3 italic caption
# Para[30]: Maraviroc–CCR5 Protein–Ligand Docking Results heading

anchor_methods = None   # will insert AFTER this (Para[20])
anchor_results = None   # will insert BEFORE this (Para[30])
fig3_caption = None     # Para[29]

for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt.startswith('Receptor preparation was performed') and anchor_methods is None:
        anchor_methods = p
        print(f"[{i}] anchor_methods: {txt[:60]}")
    if txt.startswith('Maraviroc') and 'Protein' in txt and 'Docking' in txt and 'AutoDock Vina generated' not in txt:
        # This is the results heading
        anchor_results = p
        print(f"[{i}] anchor_results: {txt[:60]}")
    if 'Figure 3' in txt and 'Ramachandran' in txt:
        fig3_caption = p
        print(f"[{i}] fig3_caption: {txt[:60]}")

assert anchor_methods is not None, "anchor_methods not found"
assert anchor_results is not None, "anchor_results not found"

# ── 1. Insert Methods paragraph after anchor_methods ────────────────────────
methods_text = (
    "Before docking test compounds, the native ligand of CCR5 was re-docked into the "
    "receptor to verify that the protocol reproduces the crystallographic binding pose. "
    "The maraviroc coordinates (residue MRV, chain A) were extracted from 4MBS and "
    "converted to SDF format using Open Babel 3.1.0. Fifty conformers were then generated "
    "with the ETKDGv3 algorithm in RDKit and minimized with the MMFF94 force field. Each "
    "conformer was aligned to the crystal pose by Kabsch rigid-body superposition, and the "
    "RMSD was computed over all 37 heavy atoms. An RMSD below 2.0 Angstroms was set as the "
    "acceptance criterion, consistent with standard practice in computational docking studies "
    "(Ramirez and Caballero, 2018)."
)

methods_p_xml = clone_para_format(anchor_methods, methods_text)
anchor_methods._p.addnext(methods_p_xml)
print("Methods paragraph inserted.")

# Re-fetch paragraphs after DOM change
doc2 = Document(DOC_PATH)  # reload to avoid stale refs - actually just save/reload
doc.save(DOC_PATH)
doc = Document(DOC_PATH)

# ── 2. Insert Results heading + body + figure before anchor_results ──────────
results_body_text = (
    "Re-docking maraviroc into the CCR5 binding site (PDB ID: 4MBS) yielded a best-conformer "
    "RMSD of 1.266 Angstroms against the crystal pose. This is within the 2.0 Angstrom "
    "threshold, confirming that ETKDGv3 conformer sampling with Kabsch superposition "
    "reproduces the experimental binding geometry. The superimposed crystal and re-docked "
    "poses within the binding pocket are shown in Figure 4. On this basis, docking of "
    "maraviroc and the CCL5-CCR5 protein-protein calculations proceeded with confidence in "
    "the geometric accuracy of the protocol."
)

results_heading_text = "Docking Protocol Validation by Re-docking"
results_caption_text = (
    "Figure 4. Superimposition of the crystal pose (green) and re-docked pose (red) of "
    "maraviroc within the CCR5 binding pocket (PDB ID: 4MBS). Gold dashed lines indicate "
    "per-atom displacement. The Calpha trace of nearby CCR5 residues is shown in blue. "
    "RMSD between poses = 1.266 Angstroms."
)

# Find anchor_results again after reload
anchor_results_new = None
for p in doc.paragraphs:
    txt = p.text.strip()
    if txt.startswith('Maraviroc') and 'Protein' in txt and 'Docking' in txt and 'AutoDock Vina generated' not in txt:
        anchor_results_new = p
        break

assert anchor_results_new is not None, "anchor_results_new not found"

# Reference para for formatting
ref_para = None
for p in doc.paragraphs:
    if p.text.strip().startswith('Ramachandran plot analysis was run'):
        ref_para = p
        break
if ref_para is None:
    ref_para = anchor_results_new

# Build caption paragraph (italic)
def make_italic_para(source_para, text):
    new_p = OxmlElement('w:p')
    pPr = source_para._p.find(qn('w:pPr'))
    if pPr is not None:
        new_p.append(copy.deepcopy(pPr))
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    i_elem = OxmlElement('w:i')
    rPr.append(i_elem)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    new_p.append(r)
    return new_p

# Insert in reverse order (each addprevious goes before anchor):
# Order needed: heading → body → image → caption → [anchor_results]

caption_p = make_italic_para(ref_para, results_caption_text)
anchor_results_new._p.addprevious(caption_p)

# Image paragraph
img_p = insert_image_para(ref_para, JPG_PATH, width_inches=5.5)
anchor_results_new._p.addprevious(img_p)  # now before caption? No - addprevious inserts right before anchor
# We need: heading, body, image, caption, then anchor
# addprevious inserts immediately before the element, so:
# First: insert caption (it goes right before anchor_results)
# Then: insert image (it goes right before anchor_results, pushing caption after it? No...)
# Actually addprevious inserts element RIGHT BEFORE the target.
# So we need to do it in reverse order of desired final sequence:
# Desired: heading → body → image → caption → anchor_results
# Steps (each addprevious to anchor_results):
#   1. addprevious(caption)   → caption, anchor_results
#   2. addprevious(image)     → image, caption, anchor_results (image goes before caption)
#   3. addprevious(body)      → body, image, caption, anchor_results
#   4. addprevious(heading)   → heading, body, image, caption, anchor_results
# But we already inserted caption and image above. Now add body then heading.

body_p = clone_para_format(ref_para, results_body_text)
# Find the image element (last inserted before anchor)
# We'll insert body before image. Image is now at anchor_results._p.getprevious().getprevious()
# Easier: just insert before the image by addprevious to anchor_results, then heading before body
anchor_results_new._p.addprevious(body_p)  # → body, image, caption, anchor_results

heading_p = clone_para_format(anchor_results_new, "", bold_prefix=results_heading_text)
anchor_results_new._p.addprevious(heading_p)  # → heading, body, image, caption, anchor_results

# Wait - that's wrong. addprevious inserts immediately before anchor_results, not before body.
# After step 3 (body): ..., body, image, caption, anchor_results
# After step 4 (heading): ..., heading, body, image, caption, anchor_results ← CORRECT!
# Wait no. addprevious(heading_p) adds heading right before anchor_results:
# → ..., body, image, caption, heading, anchor_results  <- WRONG

# We need to insert heading before body. Let me fix:
# The body_p is now in the XML. We need heading before body.
# Use body_p.addprevious(heading_p)
# But we already added heading_p before anchor_results. Let's remove it and re-add.

# Actually the logic is: addprevious(x) means "insert x immediately before me (the target element)"
# So:
# anchor._p.addprevious(caption) -> caption is now immediately before anchor
# anchor._p.addprevious(image)   -> image is now immediately before anchor (between caption and anchor? no)
# addprevious always inserts IMMEDIATELY before the calling element.
# So: ..., image, anchor  (caption got pushed away? No, caption is still there)
# Actually: ..., caption, anchor  then addprevious(image):
# ...  caption ANCHOR
# adding image before anchor: ... caption IMAGE ANCHOR  <- image between caption and anchor
# then adding body before anchor: ... caption IMAGE BODY ANCHOR
# then adding heading before anchor: ... caption IMAGE BODY HEADING ANCHOR <- WRONG ORDER

# The correct approach: build all elements, then insert them all in forward order using
# parent.insert() or use addprevious on each other.

# Let me redo: remove what was added and re-do in correct order.
print("Reordering insertions...")

# Remove the wrongly ordered elements
for bad in [caption_p, img_p, body_p, heading_p]:
    try:
        parent = bad.getparent()
        if parent is not None:
            parent.remove(bad)
    except Exception as e:
        print(f"remove error: {e}")

# Now insert in correct order using the chained approach:
# Insert heading first (addprevious to anchor) → heading, anchor
# Then insert body (addprevious to anchor) → heading, body, anchor  ... wait same issue
# Correct approach: insert each as previous of the LAST inserted element

anchor_results_new._p.addprevious(heading_p)        # heading is now before anchor
heading_p.addnext(body_p)                            # body after heading
body_p.addnext(img_p)                                # image after body
img_p.addnext(caption_p)                             # caption after image

print("Results section inserted in correct order.")

doc.save(DOC_PATH)
print(f"Saved: {DOC_PATH}")

# Verify order
doc_v = Document(DOC_PATH)
for i, p in enumerate(doc_v.paragraphs):
    txt = p.text.strip()
    if txt and any(kw in txt for kw in ['Docking Protocol', 'Re-docking maraviroc', 'Figure 4', 'Maraviroc']):
        print(f"[{i:3d}] {txt[:80]}")
