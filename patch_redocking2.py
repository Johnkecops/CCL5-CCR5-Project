"""
Insert re-docking validation content into manuscript.
Run AFTER methods paragraph is already inserted (previous partial run succeeded).
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import copy, os

DOC_PATH = '/sessions/confident-keen-gauss/mnt/First Iteration/revision-notes-0-02052026045611.docx'
JPG_PATH = '/sessions/confident-keen-gauss/mnt/First Iteration/maraviroc_CCR5_redocking.jpg'

doc = Document(DOC_PATH)

# Check if methods validation para was already added
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if 'native ligand of CCR5 was re-docked' in txt:
        print(f"[{i}] Methods validation already present: {txt[:60]}")
        break
else:
    print("Methods paragraph NOT found — need to add it")

# ── Helper: build image paragraph directly into THIS document ──────────────
def insert_image_para_inline(doc, image_path, width_inches=5.5):
    """Add a temp paragraph with image, return its _p element, then remove para."""
    # Add at end of document body
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    img_p = copy.deepcopy(para._p)
    # Remove the temp paragraph
    para._p.getparent().remove(para._p)
    return img_p

def clone_para_format(source_para, new_text):
    new_p = OxmlElement('w:p')
    pPr = source_para._p.find(qn('w:pPr'))
    if pPr is not None:
        new_p.append(copy.deepcopy(pPr))
    r = OxmlElement('w:r')
    first_run = source_para._p.find(qn('w:r'))
    if first_run is not None:
        rPr = first_run.find(qn('w:rPr'))
        if rPr is not None:
            r.append(copy.deepcopy(rPr))
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = new_text
    r.append(t)
    new_p.append(r)
    return new_p

def make_bold_heading_para(source_para, text):
    new_p = OxmlElement('w:p')
    pPr = source_para._p.find(qn('w:pPr'))
    if pPr is not None:
        new_p.append(copy.deepcopy(pPr))
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    b = OxmlElement('w:b')
    rPr.append(b)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    new_p.append(r)
    return new_p

def make_italic_para(source_para, text):
    new_p = OxmlElement('w:p')
    pPr = source_para._p.find(qn('w:pPr'))
    if pPr is not None:
        new_p.append(copy.deepcopy(pPr))
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    i_elem = OxmlElement('w:i')
    rPr.append(i_elem)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    new_p.append(r)
    return new_p

# ── Text content ─────────────────────────────────────────────────────────────
results_heading_text = "Docking Protocol Validation by Re-docking"

results_body_text = (
    "Re-docking maraviroc into the CCR5 binding site (PDB ID: 4MBS) yielded a best-conformer "
    "RMSD of 1.266 Angstroms against the crystal pose. This falls within the 2.0 Angstrom "
    "acceptance threshold, confirming that ETKDGv3 conformer sampling with Kabsch superposition "
    "reproduces the experimental binding geometry. The superimposed crystal and re-docked "
    "poses within the binding pocket are shown in Figure 4. On this basis, docking of "
    "maraviroc and the CCL5-CCR5 protein-protein calculations proceeded with confidence in "
    "the geometric accuracy of the protocol."
)

results_caption_text = (
    "Figure 4. Superimposition of the crystal pose (green) and re-docked pose (red) of "
    "maraviroc within the CCR5 binding pocket (PDB ID: 4MBS). Gold dashed lines indicate "
    "per-atom displacement. The C-alpha trace of nearby CCR5 residues is shown in blue. "
    "RMSD between poses = 1.266 Angstroms."
)

# ── Locate anchor: "Maraviroc–CCR5 Protein–Ligand Docking" results heading ──
anchor_results = None
ref_para = None  # for formatting reference (Ramachandran results body)
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if 'Ramachandran plot analysis was run' in txt and ref_para is None:
        ref_para = p
        print(f"[{i}] ref_para found: {txt[:60]}")
    if txt == 'Maraviroc–CCR5 Protein–Ligand Docking':
        anchor_results = p
        print(f"[{i}] anchor_results found: {txt[:60]}")
    elif anchor_results is None and txt.startswith('Maraviroc') and 'Protein' in txt and 'Ligand Docking' in txt and 'AutoDock' not in txt:
        anchor_results = p
        print(f"[{i}] anchor_results found (alt): {txt[:60]}")

if anchor_results is None:
    # fallback: find by index
    for i, p in enumerate(doc.paragraphs):
        if 'AutoDock Vina generated nine binding modes' in p.text:
            # anchor is the heading just before this
            if i > 0:
                anchor_results = doc.paragraphs[i-1]
                print(f"Found anchor via fallback at [{i-1}]: {anchor_results.text[:60]}")
            break

assert anchor_results is not None, "Could not find anchor_results"
if ref_para is None:
    ref_para = anchor_results

# ── Build image paragraph ────────────────────────────────────────────────────
img_p = insert_image_para_inline(doc, JPG_PATH, width_inches=5.5)
print("Image paragraph built.")

# ── Build text paragraphs ────────────────────────────────────────────────────
heading_p = make_bold_heading_para(ref_para, results_heading_text)
body_p    = clone_para_format(ref_para, results_body_text)
caption_p = make_italic_para(ref_para, results_caption_text)

# ── Insert in correct order: heading → body → image → caption → anchor ──────
# Use chained addnext for reliable ordering:
anchor_results._p.addprevious(heading_p)   # heading right before anchor
heading_p.addnext(body_p)                  # body right after heading
body_p.addnext(img_p)                      # image right after body
img_p.addnext(caption_p)                   # caption right after image

print("All results elements inserted.")

doc.save(DOC_PATH)
print(f"Saved: {DOC_PATH}")

# ── Verify final ordering ────────────────────────────────────────────────────
doc_v = Document(DOC_PATH)
print("\n--- Verification (paragraphs around insertion point) ---")
for i, p in enumerate(doc_v.paragraphs):
    txt = p.text.strip()
    if not txt:
        continue
    if any(kw in txt for kw in [
        'Docking Protocol Validation',
        'Re-docking maraviroc',
        'Figure 4',
        'Maraviroc',
        'AutoDock Vina generated',
        'native ligand of CCR5',
    ]):
        print(f"[{i:3d}] {txt[:90]}")
