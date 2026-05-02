from docx import Document

PATH = '/sessions/confident-keen-gauss/mnt/First Iteration/revision-notes-0-02052026045611.docx'
doc = Document(PATH)

# ── Para [16]: add resolution and method to both structures ───────────────────
OLD16 = ("The three-dimensional structure of CCR5 (receptor; PDB ID: 4MBS; "
         "https://www.rcsb.org/structure/4MBS) and CCL5 (ligand for "
         "protein–protein docking; PDB ID: 1RTO; "
         "https://www.rcsb.org/structure/1RTO) were retrieved from "
         "the RCSB Protein Data Bank (https://www.rcsb.org) in PDB format.")

NEW16 = ("The three-dimensional structure of CCR5 (receptor; PDB ID: 4MBS; "
         "X-ray crystallography; resolution: 2.71 Å; "
         "https://www.rcsb.org/structure/4MBS) and CCL5 (ligand for "
         "protein–protein docking; PDB ID: 1RTO; solution NMR; "
         "https://www.rcsb.org/structure/1RTO) were retrieved from "
         "the RCSB Protein Data Bank (https://www.rcsb.org) in PDB format. "
         "The 2.71 Å resolution of 4MBS is within the range accepted "
         "for reliable molecular docking, as structures resolved at "
         "≤3.0 Å are considered sufficiently accurate for binding "
         "site geometry and receptor–ligand interaction analysis.")

# ── Para [41]: add a sentence reinforcing 4MBS structural quality ─────────────
OLD41 = ("Several methodological factors may have contributed to the high RMSD "
         "in the protein–protein docking.")

NEW41 = ("The CCR5 structure used (PDB ID: 4MBS; 2.71 Å X-ray) provides "
         "a well-resolved binding pocket suitable for docking analysis; "
         "the structural quality of the receptor is a prerequisite for "
         "meaningful docking results, and the sub-3.0 Å resolution of "
         "4MBS satisfies this criterion. "
         "Several methodological factors may have contributed to the high RMSD "
         "in the protein–protein docking.")

for para_idx, old, new in [(16, OLD16, NEW16), (41, OLD41, NEW41)]:
    para = doc.paragraphs[para_idx]
    patched = False
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            print("Para[" + str(para_idx) + "] patched OK")
            patched = True
            break
    if not patched:
        print("Para[" + str(para_idx) + "] WARNING: pattern not found")

doc.save(PATH)
print("Saved.")

# Verify
doc2 = Document(PATH)
for idx in [16, 41]:
    print("\nPara[" + str(idx) + "]:\n" + doc2.paragraphs[idx].text[:500])
