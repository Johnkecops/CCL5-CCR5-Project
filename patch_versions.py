from docx import Document

PATH = '/sessions/confident-keen-gauss/mnt/First Iteration/revision-notes-0-02052026045611.docx'
doc = Document(PATH)

replacements = [
    (5,  "using AutoDock Vina, and",
         "using AutoDock Vina 1.2.7, and"),
    (14, "AutoDock Vina (Center for Computational Structural Biology)",
         "AutoDock Vina 1.2.7 (Center for Computational Structural Biology)"),
    (14, "AutoDock Tools (MGL Tools)",
         "AutoDock Tools 1.5.7 (MGL Tools)"),
    (18, "Receptor preparation was performed using AutoDock Tools:",
         "Receptor preparation was performed using AutoDock Tools 1.5.7:"),
]

for para_idx, old, new in replacements:
    para = doc.paragraphs[para_idx]
    patched = False
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            old_short = old[:55]
            new_short = new[:55]
            print("  Para[" + str(para_idx) + "] OK: '" + old_short + "' -> '" + new_short + "'")
            patched = True
            break
    if not patched:
        print("  Para[" + str(para_idx) + "] WARNING: not found: '" + old[:60] + "'")

doc.save(PATH)
print("\nSaved.")

# Verify
doc2 = Document(PATH)
for idx in [5, 14, 18]:
    t = doc2.paragraphs[idx].text
    for kw in ["Vina 1.2", "Tools 1.5"]:
        if kw in t:
            i = t.find(kw)
            print("  Para[" + str(idx) + "] verified: ..." + t[max(0,i-10):i+40] + "...")
